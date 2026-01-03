"""
学生功能路由模块
处理学生相关的所有功能：加入班级、查看作业、提交作业、查看成绩等
"""
import os
import uuid
from flask import Blueprint, jsonify, request
from werkzeug.utils import secure_filename
from db.connect import get_db
from datetime import datetime
from .utils import token_required

# 创建学生功能蓝图
student_bp = Blueprint('student', __name__)

# ==================== 配置文件上传 ====================
# 上传文件夹路径（相对于项目根目录的uploads文件夹）
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'uploads')
# 允许上传的文件扩展名（仅支持PDF和Word文档）
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# 创建上传目录（如果不存在）
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ==================== 工具函数 ====================
def allowed_file(filename):
    """
    检查文件扩展名是否允许上传
    
    Args:
        filename: 文件名
        
    Returns:
        bool: 如果文件扩展名在允许列表中返回True，否则返回False
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_file_content(file_path):
    """
    提取文件内容（支持PDF、DOCX、DOC格式）
    
    功能：
    1. 根据文件扩展名选择相应的提取方法
    2. 支持PDF（使用PyPDF2或pdfplumber）
    3. 支持DOCX（使用python-docx）
    4. 支持DOC（使用docx2txt或python-docx）
    5. 提取的文本内容用于AI批改
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容，如果失败返回错误信息
    """
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # 提取PDF内容
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ''
                    for page in pdf_reader.pages:
                        text += page.extract_text() + '\n'
                    return text.strip()
            except ImportError:
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text = ''
                        for page in pdf.pages:
                            text += page.extract_text() + '\n'
                        return text.strip()
                except ImportError:
                    return "无法提取PDF内容：请安装PyPDF2或pdfplumber库"
            except Exception as e:
                return f"PDF提取失败: {str(e)}"
        
        elif file_ext == '.docx':
            # 提取DOCX内容
            try:
                from docx import Document
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return text.strip()
            except ImportError:
                return "无法提取DOCX内容：请安装python-docx库"
            except Exception as e:
                return f"DOCX提取失败: {str(e)}"
        
        elif file_ext == '.doc':
            # 提取DOC内容（旧版Word格式）
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                return text.strip() if text else "无法提取DOC文件内容"
            except ImportError:
                try:
                    # 尝试使用python-docx（可能不支持旧格式）
                    from docx import Document
                    doc = Document(file_path)
                    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    return text.strip()
                except:
                    return "无法提取DOC内容：请安装docx2txt或python-docx库"
            except Exception as e:
                return f"DOC提取失败: {str(e)}"
        
        else:
            return f"不支持的文件格式: {file_ext}"
    
    except Exception as e:
        return f"文件提取出错: {str(e)}"


# ==================== 班级相关路由 ====================
@student_bp.route('/classes/join', methods=['POST'])
@token_required
def join_class(current_user):
    """
    学生加入班级接口
    
    功能：
    1. 验证学生身份
    2. 根据班级代码查找班级
    3. 检查是否已经加入该班级
    4. 将学生添加到班级中
    
    请求体：
        {
            "code": "班级代码（6位）"
        }
        
    Returns:
        JSON: 加入结果和班级信息
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    data = request.get_json()
    class_code = data.get('code')

    if not class_code:
        return jsonify({'success': False, 'message': '请提供班级代码'}), 400

    db = get_db()
    cursor = db.cursor()

    try:
        # 查找班级
        cursor.execute("SELECT id FROM classes WHERE code = %s", (class_code,))
        class_info = cursor.fetchone()

        if not class_info:
            return jsonify({'success': False, 'message': '班级代码不存在'}), 404

        class_id = class_info['id']

        # 检查是否已经加入
        cursor.execute(
            "SELECT id FROM student_classes WHERE student_id = %s AND class_id = %s",
            (current_user['id'], class_id)
        )
        if cursor.fetchone():
            return jsonify({'success': False, 'message': '您已经加入该班级'}), 400

        # 加入班级
        cursor.execute(
            "INSERT INTO student_classes (student_id, class_id) VALUES (%s, %s)",
            (current_user['id'], class_id)
        )
        db.commit()

        # 获取班级信息
        cursor.execute("SELECT * FROM classes WHERE id = %s", (class_id,))
        class_info = cursor.fetchone()
        cursor.close()

        return jsonify({
            'success': True,
            'message': '成功加入班级',
            'class': class_info
        })
    except Exception as e:
        db.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        cursor.close()


@student_bp.route('/classes', methods=['GET'])
@token_required
def get_student_classes(current_user):
    """
    获取学生已加入的班级列表
    
    功能：
    1. 查询学生加入的所有班级
    2. 返回班级详细信息（名称、代码、描述、加入时间等）
    
    Returns:
        JSON: 班级列表
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    db = get_db()
    cursor = db.cursor()

    cursor.execute(
        """SELECT c.*, sc.joined_at 
           FROM student_classes sc 
           JOIN classes c ON sc.class_id = c.id 
           WHERE sc.student_id = %s 
           ORDER BY sc.joined_at DESC""",
        (current_user['id'],)
    )
    classes = cursor.fetchall()
    cursor.close()

    return jsonify(classes)


# ==================== 作业相关路由 ====================
@student_bp.route('/assignments', methods=['GET'])
@token_required
def get_student_assignments(current_user):
    """
    获取学生可见的作业列表
    
    功能：
    1. 查询学生已加入的所有班级
    2. 获取这些班级的所有作业
    3. 标记每个作业是否已提交
    4. 只显示学生已加入班级的作业（权限控制）
    
    Returns:
        JSON: 作业列表，每个作业包含是否已提交的标记
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    db = get_db()
    cursor = db.cursor()

    # 获取学生已加入的班级ID列表
    cursor.execute(
        "SELECT class_id FROM student_classes WHERE student_id = %s",
        (current_user['id'],)
    )
    class_ids = [row['class_id'] for row in cursor.fetchall()]

    # 如果没有加入任何班级，返回空列表
    if not class_ids:
        cursor.close()
        return jsonify([])

    # 获取这些班级的作业
    placeholders = ','.join(['%s'] * len(class_ids))
    cursor.execute(
        f"SELECT * FROM assignments WHERE class_id IN ({placeholders}) ORDER BY deadline ASC",
        tuple(class_ids)
    )
    assignments = cursor.fetchall()

    # 获取学生已提交的作业ID
    cursor.execute(
        "SELECT assignment_id FROM submissions WHERE student_id = %s",
        (current_user['id'],)
    )
    submitted_ids = [item['assignment_id'] for item in cursor.fetchall()]

    cursor.close()

    # 添加是否已提交的标记
    for assignment in assignments:
        assignment['submitted'] = assignment['id'] in submitted_ids

    return jsonify(assignments)


@student_bp.route('/assignments/<int:assignment_id>/submit', methods=['POST'])
@token_required
def submit_assignment(current_user, assignment_id):
    """
    学生提交作业接口
    
    功能：
    1. 接收文本内容和/或文件上传
    2. 保存文件到服务器
    3. 提取文件内容（用于AI批改）
    4. 调用AI进行自动批改
    5. 保存提交记录到数据库
    6. 如果已提交过，则更新提交记录
    
    请求方式：
        POST (multipart/form-data)
        
    请求参数：
        - content: 文本内容（可选）
        - file: 文件（可选，支持PDF、DOC、DOCX）
        
    Args:
        assignment_id: 作业ID
        
    Returns:
        JSON: 提交结果，包含AI评分和评价
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    db = get_db()
    cursor = db.cursor()

    try:
        # 获取文本内容
        content = request.form.get('content', '')

        # 处理文件上传
        file_url = ''
        file_content = ''
        if 'file' in request.files:
            file = request.files['file']
            if file.filename != '' and allowed_file(file.filename):
                # 生成唯一文件名（避免冲突）
                filename = f"{uuid.uuid4()}_{secure_filename(file.filename)}"
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                file.save(file_path)

                # 生成可访问的URL
                file_url = f"/uploads/{filename}"
                
                # 提取文件内容用于AI批改
                try:
                    extracted = extract_file_content(file_path)
                    # 检查是否是错误消息
                    if extracted and not (extracted.startswith("无法提取") or 
                                         extracted.startswith("提取失败") or 
                                         extracted.startswith("不支持") or
                                         extracted.startswith("文件提取出错")):
                        file_content = extracted
                    else:
                        print(f"文件内容提取警告: {extracted}")
                        file_content = ''  # 提取失败时设为空
                except Exception as e:
                    print(f"文件内容提取出错: {str(e)}")
                    file_content = ''

        # 获取作业标题，用于AI批改
        cursor.execute("SELECT title, description FROM assignments WHERE id = %s", (assignment_id,))
        assignment = cursor.fetchone()
        assignment_description = assignment['description'] if assignment else "未知作业"

        # AI自动批改（支持文本内容和文件内容）
        ai_score = None
        ai_feedback = None
        
        # 组合文本内容和文件内容
        submission_content = ''
        if content and content.strip():
            submission_content = content.strip()
        if file_content and file_content.strip():
            if submission_content:
                submission_content += "\n\n--- 文件内容 ---\n\n" + file_content.strip()
            else:
                submission_content = file_content.strip()
        
        # 如果有提交内容（文本或文件），则进行AI批改
        if submission_content:
            try:
                from ai_grading import TongyiQianwenAPI
                ai = TongyiQianwenAPI()
                ai_score, ai_feedback = ai.get_grading(assignment_description, submission_content)
            except Exception as e:
                print(f"AI批改服务初始化失败: {str(e)}")
                ai_feedback = "AI批改服务暂时不可用"
        else:
            ai_feedback = "未提供文本内容或文件，无法进行AI批改"

        # 检查是否已提交
        cursor.execute(
            "SELECT id FROM submissions WHERE assignment_id = %s AND student_id = %s",
            (assignment_id, current_user['id'])
        )
        if cursor.fetchone():
            # 更新提交
            cursor.execute(
                """UPDATE submissions SET content = %s, file_url = %s, submitted_at = %s,
                   ai_score = %s, ai_feedback = %s
                   WHERE assignment_id = %s AND student_id = %s""",
                (content, file_url, datetime.now(), ai_score, ai_feedback,
                 assignment_id, current_user['id'])
            )
        else:
            # 新提交
            cursor.execute(
                """INSERT INTO submissions 
                   (assignment_id, student_id, content, file_url, submitted_at,
                    ai_score, ai_feedback)
                   VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                (assignment_id, current_user['id'], content, file_url, datetime.now(),
                 ai_score, ai_feedback)
            )

        db.commit()
        return jsonify({
            'success': True,
            'message': '作业提交成功',
            'file_url': file_url,
            'ai_score': ai_score,
            'ai_feedback': ai_feedback
        })
    except Exception as e:
        db.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        cursor.close()


# ==================== 提交历史相关路由 ====================
@student_bp.route('/submissions', methods=['GET'])
@token_required
def get_student_submissions(current_user):
    """
    获取学生的提交历史
    
    功能：
    1. 查询学生所有已提交的作业
    2. 包含作业标题、提交时间、成绩等信息
    3. 按提交时间倒序排列（最新的在前）
    
    Returns:
        JSON: 提交历史列表
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    db = get_db()
    cursor = db.cursor()

    cursor.execute(
        """SELECT s.*, a.title FROM submissions s
           JOIN assignments a ON s.assignment_id = a.id
           WHERE s.student_id = %s ORDER BY submitted_at DESC""",
        (current_user['id'],)
    )
    submissions = cursor.fetchall()
    cursor.close()

    return jsonify(submissions)


# ==================== 成绩相关路由 ====================
@student_bp.route('/grades/trend', methods=['GET'])
@token_required
def get_grade_trend(current_user):
    """
    获取学生各班级最近5次提交的成绩数据（用于趋势分析）
    
    功能：
    1. 查询学生已加入的所有班级
    2. 对每个班级，获取最近5次已批改的提交
    3. 优先使用教师批改的分数，如果没有则使用AI评分
    4. 按时间正序排列（用于绘制趋势图）
    
    Returns:
        JSON: 按班级分组的成绩数据
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    db = get_db()
    cursor = db.cursor()

    try:
        # 获取学生已加入的班级
        cursor.execute(
            """SELECT c.id, c.name, c.code 
               FROM student_classes sc 
               JOIN classes c ON sc.class_id = c.id 
               WHERE sc.student_id = %s 
               ORDER BY c.name""",
            (current_user['id'],)
        )
        classes = cursor.fetchall()

        result = []
        
        for class_info in classes:
            class_id = class_info['id']
            class_name = class_info['name']
            
            # 获取该班级最近5次已批改的提交（优先使用教师批改的score，如果没有则使用ai_score）
            cursor.execute(
                """SELECT s.id, s.submitted_at, 
                          COALESCE(s.score, s.ai_score) as grade,
                          a.title as assignment_title
                   FROM submissions s
                   JOIN assignments a ON s.assignment_id = a.id
                   WHERE s.student_id = %s 
                     AND a.class_id = %s
                     AND (s.score IS NOT NULL OR s.ai_score IS NOT NULL)
                   ORDER BY s.submitted_at DESC
                   LIMIT 5""",
                (current_user['id'], class_id)
            )
            submissions = cursor.fetchall()
            
            # 如果该班级有成绩数据，添加到结果中
            if submissions:
                # 按时间正序排列（最早的在前）
                submissions_sorted = sorted(submissions, key=lambda x: x['submitted_at'])
                
                result.append({
                    'class_id': class_id,
                    'class_name': class_name,
                    'class_code': class_info['code'],
                    'submissions': [
                        {
                            'id': s['id'],
                            'submitted_at': s['submitted_at'],
                            'grade': s['grade'],
                            'assignment_title': s['assignment_title']
                        }
                        for s in submissions_sorted
                    ]
                })
        
        cursor.close()
        return jsonify({
            'success': True,
            'data': result
        })
    
    except Exception as e:
        cursor.close()
        return jsonify({
            'success': False,
            'message': f'获取成绩趋势失败：{str(e)}'
        }), 500


@student_bp.route('/grades/trend/analyze', methods=['POST'])
@token_required
def analyze_grade_trend(current_user):
    """
    获取AI成绩趋势分析
    
    功能：
    1. 接收班级名称和成绩数据
    2. 调用AI分析成绩趋势
    3. 返回分析结果（趋势、原因、建议等）
    
    请求体：
        {
            "class_name": "班级名称",
            "grade_data": [
                {
                    "submitted_at": "提交时间",
                    "grade": 分数,
                    "assignment_title": "作业标题"
                }, ...
            ]
        }
        
    Returns:
        JSON: AI分析结果
    """
    if current_user['role'] != 'student':
        return jsonify({'message': '权限不足'}), 403

    data = request.get_json()
    class_name = data.get('class_name')
    grade_data = data.get('grade_data', [])

    if not class_name or not grade_data:
        return jsonify({
            'success': False,
            'message': '缺少必要参数'
        }), 400

    try:
        from ai_grading import TongyiQianwenAPI
        ai = TongyiQianwenAPI()
        analysis, success = ai.analyze_grade_trend(class_name, grade_data)
        
        return jsonify({
            'success': success,
            'analysis': analysis
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'AI分析失败：{str(e)}'
        }), 500
